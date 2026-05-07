"""
Tool: generar_kit
Genera el kit de entrevista en .docx usando python-docx.
Sin dependencias de Node.js ni npm.
"""

import os
import shutil
import subprocess
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_OUTPUT_DIR = Path(os.environ.get("KIT_OUTPUT_DIR", "./output/kits"))
_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

PURPLE_DARK  = RGBColor(0x4A, 0x3B, 0x8C)
PURPLE_MID   = RGBColor(0x6B, 0x56, 0xC8)
TEAL_DARK    = RGBColor(0x0F, 0x6E, 0x56)
CORAL_DARK   = RGBColor(0x99, 0x3C, 0x1D)
GRAY_DARK    = RGBColor(0x44, 0x44, 0x41)
GRAY_MID     = RGBColor(0x88, 0x87, 0x80)
GRAY_LIGHT   = RGBColor(0xF1, 0xEF, 0xE8)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

CATEGORIA_CONFIG = {
    "técnica":    {"fill": "EAE7F8", "text": PURPLE_DARK,  "label": "TÉCNICA"},
    "conductual": {"fill": "E1F5EE", "text": TEAL_DARK,    "label": "CONDUCTUAL"},
    "presión":    {"fill": "FAECE7", "text": CORAL_DARK,   "label": "PRESIÓN / DETECCIÓN"},
    "cultura":    {"fill": "FAEEDA", "text": RGBColor(0x85, 0x4F, 0x0B), "label": "CULTURA"},
}

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_bottom_border(para, color="D3D1C7"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def _para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spac = OxmlElement("w:spacing")
    spac.set(qn("w:before"), str(before))
    spac.set(qn("w:after"), str(after))
    pPr.append(spac)

def _run(para, text, bold=False, italic=False, size=10, color=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def _hr(doc, color="6B56C8"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    _para_spacing(p, before=0, after=100)

def _page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def _build_portada(doc, data):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "KIT DE ENTREVISTA", bold=True, size=24, color=PURPLE_DARK)
    _para_spacing(p, before=300, after=80)
    _hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, data["candidato_nombre"], bold=True, size=18, color=GRAY_DARK)
    _para_spacing(p, before=80, after=40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, data.get("proceso_titulo") or "Proceso de selección", size=12, color=GRAY_MID)
    _para_spacing(p, before=0, after=160)
    fecha = datetime.now().strftime("%d/%m/%Y")
    rows_data = [
        ("Candidato",     data["candidato_nombre"]),
        ("Proceso",       data.get("proceso_titulo") or "—"),
        ("Preguntas",     f"{data.get('preguntas_total', 0)} preguntas"),
        ("Duración est.", f"{data['duracion_estimada_min']} min" if data.get("duracion_estimada_min") else "—"),
        ("Fecha",         fecha),
    ]
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        _set_cell_bg(t.rows[i].cells[0], "F1EFE8")
        _run(t.rows[i].cells[0].paragraphs[0], label, bold=True, size=9, color=GRAY_DARK)
        _run(t.rows[i].cells[1].paragraphs[0], value, size=9, color=GRAY_DARK)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Documento confidencial · Uso interno de RRHH", italic=True, size=8, color=GRAY_MID)
    _page_break(doc)

def _build_perfil(doc, data):
    h = doc.add_heading("Perfil del candidato", level=1)
    h.runs[0].font.color.rgb = PURPLE_DARK
    _hr(doc)
    skills = data.get("skills", [])
    if skills:
        h2 = doc.add_heading("Skills declarados", level=2)
        h2.runs[0].font.color.rgb = GRAY_DARK
        chunks = [skills[i:i+4] for i in range(0, len(skills), 4)]
        for chunk in chunks:
            chunk += [""] * (4 - len(chunk))
            t = doc.add_table(rows=1, cols=4)
            t.style = "Table Grid"
            for j, skill in enumerate(chunk):
                cell = t.rows[0].cells[j]
                if skill:
                    _set_cell_bg(cell, "EAE7F8")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _run(p, skill, bold=bool(skill), size=9, color=PURPLE_DARK)
            doc.add_paragraph()
    experiencia = data.get("experiencia", [])
    if experiencia:
        h2 = doc.add_heading("Experiencia laboral", level=2)
        h2.runs[0].font.color.rgb = GRAY_DARK
        for exp in experiencia:
            p = doc.add_paragraph()
            _run(p, exp.get("cargo", ""), bold=True, size=11, color=GRAY_DARK)
            _run(p, f"  ·  {exp.get('empresa', '')}", size=11, color=GRAY_MID)
            _para_spacing(p, before=80, after=20)
            p2 = doc.add_paragraph()
            _run(p2, f"{exp.get('desde','?')} → {exp.get('hasta') or 'actual'}", italic=True, size=9, color=GRAY_MID)
            _para_spacing(p2, before=0, after=20)
            if exp.get("descripcion"):
                p3 = doc.add_paragraph()
                _run(p3, exp["descripcion"], size=9, color=GRAY_DARK)
                _para_spacing(p3, before=0, after=60)
    _page_break(doc)

def _build_preguntas(doc, preguntas):
    h = doc.add_heading("Preguntas de entrevista", level=1)
    h.runs[0].font.color.rgb = PURPLE_DARK
    _hr(doc)
    grupos = {}
    for p in preguntas:
        cat = p.get("categoria", "técnica")
        grupos.setdefault(cat, []).append(p)
    num = 1
    for cat, prgs in grupos.items():
        cfg = CATEGORIA_CONFIG.get(cat, CATEGORIA_CONFIG["técnica"])
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        _set_cell_bg(t.rows[0].cells[0], cfg["fill"])
        _run(t.rows[0].cells[0].paragraphs[0], f"● {cfg['label']}  ({len(prgs)} preguntas)", bold=True, size=10, color=cfg["text"])
        doc.add_paragraph()
        for pregunta in prgs:
            t2 = doc.add_table(rows=1, cols=2)
            t2.style = "Table Grid"
            cn = t2.rows[0].cells[0]
            cn.width = Cm(1.2)
            _set_cell_bg(cn, cfg["fill"])
            pn = cn.paragraphs[0]
            pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(pn, str(num), bold=True, size=13, color=cfg["text"])
            cc = t2.rows[0].cells[1]
            pq = cc.add_paragraph()
            _run(pq, pregunta.get("pregunta", ""), bold=True, size=10, color=GRAY_DARK)
            _para_spacing(pq, before=0, after=40)
            po = cc.add_paragraph()
            _run(po, "Objetivo: ", bold=True, size=8, color=GRAY_MID)
            _run(po, pregunta.get("objetivo", "—"), italic=True, size=8, color=GRAY_MID)
            _para_spacing(po, before=0, after=30)
            pt = cc.add_paragraph()
            _run(pt, f"⏱ {pregunta.get('tiempo_estimado_min', 3)} min", size=8, color=GRAY_MID)
            _para_spacing(pt, before=0, after=50)
            pn2 = cc.add_paragraph()
            _run(pn2, "Notas:", bold=True, size=8, color=GRAY_MID)
            for _ in range(3):
                pl = cc.add_paragraph()
                _set_bottom_border(pl)
                _para_spacing(pl, before=0, after=120)
            doc.add_paragraph()
            num += 1

def _build_scorecard(doc):
    _page_break(doc)
    h = doc.add_heading("Scorecard de evaluación", level=1)
    h.runs[0].font.color.rgb = PURPLE_DARK
    _hr(doc)
    p = doc.add_paragraph()
    _run(p, "Completar durante o después de la entrevista.", italic=True, size=9, color=GRAY_MID)
    _para_spacing(p, before=0, after=120)
    categorias = [
        ("Conocimiento técnico",                 "30%"),
        ("Resolución de problemas",              "20%"),
        ("Comunicación",                         "15%"),
        ("Experiencia declarada vs. demostrada", "20%"),
        ("Cultura y valores",                    "15%"),
    ]
    headers = ["Criterio", "Peso", "1–3", "4–6", "7–10"]
    t = doc.add_table(rows=1 + len(categorias), cols=5)
    t.style = "Table Grid"
    for j, h_txt in enumerate(headers):
        _set_cell_bg(t.rows[0].cells[j], "4A3B8C")
        ph = t.rows[0].cells[j].paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(ph, h_txt, bold=True, size=9, color=WHITE)
    for i, (label, peso) in enumerate(categorias):
        fill = "FFFFFF" if i % 2 == 0 else "F1EFE8"
        row = t.rows[i + 1]
        _set_cell_bg(row.cells[0], fill)
        _run(row.cells[0].paragraphs[0], label, size=9, color=GRAY_DARK)
        _set_cell_bg(row.cells[1], fill)
        pp = row.cells[1].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(pp, peso, bold=True, size=9, color=GRAY_MID)
        for j in range(2, 5):
            _set_cell_bg(row.cells[j], fill)
    doc.add_paragraph()
    h2 = doc.add_heading("Decisión final", level=2)
    h2.runs[0].font.color.rgb = GRAY_DARK
    decisiones = ["Avanzar", "Avanzar con reservas", "En espera", "Descartar"]
    td = doc.add_table(rows=1, cols=4)
    td.style = "Table Grid"
    for j, dec in enumerate(decisiones):
        pd = td.rows[0].cells[j].paragraphs[0]
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(pd, f"☐  {dec}", size=9, color=GRAY_DARK)
    doc.add_paragraph()
    pc = doc.add_paragraph()
    _run(pc, "Comentarios finales:", bold=True, size=9, color=GRAY_DARK)
    _para_spacing(pc, before=80, after=40)
    for _ in range(4):
        pl = doc.add_paragraph()
        _set_bottom_border(pl)
        _para_spacing(pl, before=0, after=180)
    pf = doc.add_paragraph()
    _run(pf, "Entrevistador: ", bold=True, size=9, color=GRAY_DARK)
    _run(pf, "________________________________    ", size=9)
    _run(pf, "Fecha: ", bold=True, size=9, color=GRAY_DARK)
    _run(pf, "________________", size=9)


def generar_kit(
    candidato_id: str,
    candidato_nombre: str,
    proceso_id: str,
    proceso_titulo: str,
    skills: list[str],
    experiencia: list[dict],
    preguntas: list[dict],
    duracion_estimada_min: int | None = None,
    generar_pdf: bool = False,
) -> dict:
    """
    Genera el kit de entrevista completo en formato .docx usando python-docx.
    No requiere Node.js ni npm — funciona en cualquier entorno Python.

    Args:
        candidato_id: UUID del candidato.
        candidato_nombre: Nombre completo del candidato.
        proceso_id: UUID del proceso de selección.
        proceso_titulo: Título del proceso/posición.
        skills: Lista de skills declarados.
        experiencia: Lista de dicts con campos: empresa, cargo, desde, hasta, descripcion.
        preguntas: Lista de dicts con campos: categoria, pregunta, objetivo, tiempo_estimado_min.
        duracion_estimada_min: Duración estimada de la entrevista en minutos.
        generar_pdf: Si True, intenta convertir a PDF con LibreOffice.

    Returns:
        Dict con rutas al .docx/.pdf y metadata.
    """
    data = {
        "candidato_id":          candidato_id,
        "candidato_nombre":      candidato_nombre,
        "proceso_id":            proceso_id,
        "proceso_titulo":        proceso_titulo,
        "skills":                skills or [],
        "experiencia":           experiencia or [],
        "preguntas_total":       len(preguntas or []),
        "duracion_estimada_min": duracion_estimada_min,
    }
    nombre_safe = candidato_nombre.replace(" ", "_").lower()[:30]
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M")
    filename    = f"kit_{nombre_safe}_{timestamp}"
    docx_path   = _OUTPUT_DIR / f"{filename}.docx"
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)
        _build_portada(doc, data)
        _build_perfil(doc, data)
        _build_preguntas(doc, preguntas or [])
        _build_scorecard(doc)
        doc.save(str(docx_path))
    except Exception as e:
        return {"error": f"Error generando .docx: {e}", "kit_path_docx": None, "kit_path_pdf": None}
    if not docx_path.exists() or docx_path.stat().st_size < 1000:
        return {"error": "Archivo .docx inválido.", "kit_path_docx": None, "kit_path_pdf": None}
    pdf_path = _convertir_pdf(docx_path) if generar_pdf else None
    return {
        "kit_path_docx":         str(docx_path),
        "kit_path_pdf":          str(pdf_path) if pdf_path else None,
        "preguntas_total":       len(preguntas or []),
        "duracion_estimada_min": duracion_estimada_min,
        "filename":              filename,
    }

def _convertir_pdf(docx_path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    try:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)], capture_output=True, timeout=60)
        p = docx_path.with_suffix(".pdf")
        return p if p.exists() else None
    except Exception:
        return None
