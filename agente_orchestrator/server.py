"""
Servidor HTTP del Orchestrator de Reclutamiento.

Expone los endpoints que el frontend PWA espera:
  POST /chat              → inicia un mensaje, devuelve request_id
  GET  /chat/stream/{id}  → SSE con pasos intermedios + respuesta final
  GET  /chat/status/{id}  → polling acumulativo, alternativa a SSE

Un request_id se consume por un solo transporte: el primero que lo tome (SSE o
polling) dispara la corrida del agente; el otro devuelve 404.

Cómo correr:
    python server.py
"""

import asyncio
import base64
import json
import logging
import os
import sys
import time
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import AsyncGenerator, Optional

from dotenv import load_dotenv

load_dotenv(override=True)

from observability import init_observability  # noqa: E402 — must precede google.adk
init_observability("recruiting_orchestrator")

import uvicorn  # noqa: E402
from fastapi import FastAPI, HTTPException  # noqa: E402
from fastapi.middleware.cors import CORSMiddleware  # noqa: E402
from fastapi.responses import StreamingResponse  # noqa: E402
from google.adk.runners import Runner  # noqa: E402
from google.adk.sessions import InMemorySessionService  # noqa: E402
from google.genai import types  # noqa: E402
from opentelemetry import trace  # noqa: E402
from pydantic import BaseModel  # noqa: E402

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

sys.path.insert(0, ".")

from agent import root_agent  # noqa: E402
from request_state import (  # noqa: E402
    DEFAULT_TTL_SECONDS,
    NO_FINAL_RESPONSE_CODE,
    NO_FINAL_RESPONSE_MESSAGE,
    RequestStateStore,
)

_tracer = trace.get_tracer(__name__)

APP_NAME = "recruiting_orchestrator"
USER_ID  = "frontend_user"

# Mensaje genérico para el cliente: el detalle del error va al log, no a la respuesta.
ORCHESTRATOR_ERROR_MESSAGE = "El orchestrator no pudo completar el pedido. Reintentá en unos segundos."

# Cota para una corrida en background (polling). Sin esto el cliente pollea para siempre.
RUN_TIMEOUT_SECONDS = int(os.getenv("RUN_TIMEOUT_SECONDS", "600"))
RUN_TIMEOUT_MESSAGE = "El orchestrator tardó demasiado en responder. Reintentá."

# Cuánto sobrevive un request_id que nadie reclamó por SSE ni por polling.
PENDING_TTL_SECONDS = int(os.getenv("PENDING_TTL_SECONDS", "900"))

session_service: InMemorySessionService
runner: Runner

pending_requests: dict[str, dict] = {}
conversation_sessions: dict[str, str] = {}

request_states = RequestStateStore(
    ttl_seconds=int(os.getenv("REQUEST_STATE_TTL_SECONDS", str(DEFAULT_TTL_SECONDS)))
)

# El event loop sólo guarda referencias débiles a las tasks: sin esto, una corrida
# en background puede ser recolectada a mitad de camino.
_background_tasks: set[asyncio.Task] = set()


@asynccontextmanager
async def lifespan(app: FastAPI):
    global session_service, runner
    session_service = InMemorySessionService()
    runner = Runner(
        agent=root_agent,
        app_name=APP_NAME,
        session_service=session_service,
    )
    logger.info("Orchestrator listo → http://%s:%s", os.environ["HOST"], os.environ["PORT"])
    yield


app = FastAPI(title="Recruiting Orchestrator", version="1.0.0", lifespan=lifespan)

_cors_origins = [o.strip() for o in os.environ["CORS_ALLOWED_ORIGINS"].split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class FileAttachment(BaseModel):
    base64:   str
    fileName: str
    mimeType: str


class ChatRequest(BaseModel):
    conversation_id: Optional[str] = None
    message:         str
    file:            Optional[FileAttachment] = None  # CV adjunto (PDF/Word)
    files:           Optional[list[FileAttachment]] = None


class ChatInitResponse(BaseModel):
    conversation_id: str
    request_id:      str
    stream_url:      str


class AgentStepModel(BaseModel):
    agent:     str
    status:    str
    message:   str
    timestamp: str


class FinalMessageModel(BaseModel):
    role:    str
    content: str


class StreamErrorModel(BaseModel):
    code:    str
    message: str


class ChatStatusResponse(BaseModel):
    """Snapshot acumulativo de un request; el cliente descarta los pasos ya vistos."""
    status: str  # running | done | error
    steps:  list[AgentStepModel]
    final:  Optional[FinalMessageModel] = None
    error:  Optional[StreamErrorModel]  = None


# ---------------------------------------------------------------------------
# Eventos (compartidos por SSE y polling)
# ---------------------------------------------------------------------------

def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def _step_event(agent: str, status: str, message: str) -> dict:
    return {
        "type": "step",
        "data": {"agent": agent, "status": status, "message": message, "timestamp": _now_iso()},
    }

def _final_event(content: str) -> dict:
    return {"type": "final", "data": {"role": "assistant", "content": content}}

def _error_event(code: str, message: str) -> dict:
    return {"type": "error", "data": {"code": code, "message": message}}

def _sse_frame(event: dict) -> str:
    return f"event: {event['type']}\ndata: {json.dumps(event['data'], ensure_ascii=False)}\n\n"

def _extract_step_from_function_call(part) -> tuple[str, str] | None:
    fc = getattr(part, "function_call", None)
    if not fc:
        return None
    args = dict(fc.args) if fc.args else {}
    agent_name = args.get("agent_name", fc.name)
    payload    = args.get("payload", {})
    action     = payload.get("action", "procesando") if isinstance(payload, dict) else "procesando"
    return agent_name, action

def _extract_step_from_function_response(part) -> tuple[str, str, str] | None:
    fr = getattr(part, "function_response", None)
    if not fr:
        return None
    response   = dict(fr.response) if fr.response else {}
    raw_status = response.get("status", "done")
    step_status = "error" if raw_status == "error" else "done"
    msg = response.get("message", response.get("mensaje", f"{fr.name} completado"))
    if not isinstance(msg, str):
        msg = json.dumps(msg, ensure_ascii=False)[:200]
    return fr.name, step_status, msg


def _build_adk_content(message: str, file: Optional[FileAttachment], files: Optional[list[FileAttachment]] = None) -> types.Content:
    # Múltiples CVs para rankeo
    if files and len(files) > 0:
        logger.info("BUILD ADK | procesando %d archivos via files", len(files))
        textos = []
        for f in files:
            texto = _extraer_texto_base64(f.base64, f.mimeType)
            if texto:
                textos.append(f"=== CV: {f.fileName} ===\n{texto[:2000]}")
            else:
                textos.append(f"=== CV: {f.fileName} === [No se pudo extraer texto]")
        cvs_block = "\n\n".join(textos)
        text = f"{message}\n\n=== CVs ADJUNTOS PARA RANKEO ({len(files)} archivos) ===\n{cvs_block}"
        return types.Content(role="user", parts=[types.Part(text=text)])

    # Un solo CV
    if file:
        logger.info("BUILD ADK | procesando archivo via file")
        texto_cv = _extraer_texto_base64(file.base64, file.mimeType)
        if texto_cv:
            text = f"{message}\n\n=== CV ADJUNTO ({file.fileName}) ===\n{texto_cv[:3000]}"
        else:
            text = f"{message}\n\nNota: Se adjuntó el archivo {file.fileName} pero no se pudo extraer el texto."
        return types.Content(role="user", parts=[types.Part(text=text)])
    logger.info("BUILD ADK | solo texto, sin archivo")
    return types.Content(role="user", parts=[types.Part(text=message)])

def _extraer_texto_base64(cv_base64: str, mime_type: str) -> str | None:
    import base64
    try:
        file_bytes = base64.b64decode(cv_base64)
        logger.info("Extrayendo CV | mime=%s | base64_len=%s", mime_type, len(cv_base64) if cv_base64 else 0)
        if "pdf" in mime_type:
            import fitz
            doc  = fitz.open(stream=file_bytes, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text.strip() or None
        elif "word" in mime_type or "openxmlformats" in mime_type:
            import io
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))

            text_parts = []
            seen = set()

            for p in doc.paragraphs:
                t = p.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    text_parts.append(t)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            t = p.text.strip()
                            if t and t not in seen:
                                seen.add(t)
                                text_parts.append(t)

            text = "\n".join(text_parts)
            return text.strip() or None
    except Exception as e:
        logger.error("Error extrayendo texto: %s", e)
    return None

# ---------------------------------------------------------------------------
# POST /chat
# ---------------------------------------------------------------------------

@app.post("/chat", response_model=ChatInitResponse)
async def post_chat(req: ChatRequest) -> ChatInitResponse:
    request_states.purge_expired()
    _purge_pending()

    request_id      = str(uuid.uuid4())
    conversation_id = req.conversation_id or str(uuid.uuid4())

    if conversation_id not in conversation_sessions:
        session_id = f"session_{conversation_id}"
        conversation_sessions[conversation_id] = session_id
    else:
        session_id = conversation_sessions[conversation_id]

    pending_requests[request_id] = {
        "conversation_id": conversation_id,
        "session_id":      session_id,
        "message":         req.message,
        "file":            req.file,
        "files":           req.files,
        "created_at":      time.monotonic(),
    }

    logger.info("POST /chat → request_id=%s conv=%s file=%s files=%d",
            request_id, conversation_id,
            req.file.fileName if req.file else "none",
            len(req.files) if req.files else 0)

    return ChatInitResponse(
        conversation_id=conversation_id,
        request_id=request_id,
        stream_url=f"/chat/stream/{request_id}",
    )


# ---------------------------------------------------------------------------
# Corrida del agente — fuente única de eventos para SSE y polling
# ---------------------------------------------------------------------------

async def _run_agent_events(request_id: str, req_data: dict) -> AsyncGenerator[dict, None]:
    """Corre el agente y va emitiendo eventos (`step`, `final`, `error`).

    Args:
        request_id: Identificador del request, sólo para trazas.
        req_data: Datos guardados por POST /chat (session_id, message, file, files).

    Yields:
        Diccionarios con la forma {"type": ..., "data": {...}}.
    """
    session_id = req_data["session_id"]
    message    = req_data["message"]
    file       = req_data.get("file")
    files      = req_data.get("files")

    try:
        await session_service.create_session(
            app_name=APP_NAME, user_id=USER_ID, session_id=session_id
        )
    except Exception:
        pass

    content = _build_adk_content(message, file, files)

    # Todo camino de salida tiene que emitir un evento terminal (`final` o
    # `error`). Si el generador termina sin ninguno, el StreamingResponse cierra
    # limpio y —por spec de EventSource— el browser reconecta solo a la misma
    # URL; pero `_take_pending` ya consumió el request_id, así que el reintento
    # se come un 404 y el usuario ve un "conexión cerrada" espurio en lugar del
    # motivo real. Esto espeja la garantía que `RequestStateStore.finish()` ya
    # da en el transporte de polling.
    emitio_terminal = False

    with _tracer.start_as_current_span(f"agent.{APP_NAME}") as span:
        span.set_attribute("input.value", message[:2000])
        try:
            async for event in runner.run_async(
                user_id=USER_ID, session_id=session_id, new_message=content
            ):
                if not event.content or not event.content.parts:
                    continue

                for part in event.content.parts:
                    result = _extract_step_from_function_call(part)
                    if result:
                        agent_name, action = result
                        logger.info("Tool call → %s (%s)", agent_name, action)
                        yield _step_event(agent_name, "running", f"Consultando {agent_name}: {action}")

                    result = _extract_step_from_function_response(part)
                    if result:
                        agent_name, step_status, msg = result
                        logger.info("Tool response ← %s (%s): %s", agent_name, step_status, msg[:80])
                        yield _step_event(agent_name, step_status, msg)

                if event.is_final_response():
                    text = next(
                        (p.text for p in event.content.parts if getattr(p, "text", None)),
                        None,
                    )
                    if text:
                        logger.info("Respuesta final (%d chars)", len(text))
                        span.set_attribute("output.value", text[:2000])
                        emitio_terminal = True
                        yield _final_event(text)

        except asyncio.CancelledError:
            logger.info("Corrida cancelada por el cliente (request_id=%s)", request_id)
            raise
        except Exception as exc:
            logger.error("Error corriendo el agente (request_id=%s): %s", request_id, exc, exc_info=True)
            emitio_terminal = True
            yield _error_event("ORCHESTRATOR_ERROR", ORCHESTRATOR_ERROR_MESSAGE)

        if not emitio_terminal:
            logger.warning(
                "La corrida terminó sin respuesta final (request_id=%s)", request_id
            )
            yield _error_event(NO_FINAL_RESPONSE_CODE, NO_FINAL_RESPONSE_MESSAGE)


def _take_pending(request_id: str) -> dict:
    """Reclama un request_id pendiente. Un request_id se consume una sola vez."""
    req_data = pending_requests.pop(request_id, None)
    if not req_data:
        raise HTTPException(status_code=404, detail="request_id no encontrado o ya consumido")
    return req_data


def _purge_pending() -> int:
    """Descarta requests que nunca fueron reclamados por ningún transporte.

    Cada entrada retiene el archivo en base64, así que dejarlas acumular filtra
    CVs enteros en memoria por toda la vida del proceso.
    """
    now = time.monotonic()
    stale = [
        request_id
        for request_id, data in pending_requests.items()
        if now - data["created_at"] > PENDING_TTL_SECONDS
    ]
    for request_id in stale:
        del pending_requests[request_id]
    if stale:
        logger.info("Purgados %d requests pendientes nunca reclamados", len(stale))
    return len(stale)


# ---------------------------------------------------------------------------
# GET /chat/stream/{request_id}  — SSE
# ---------------------------------------------------------------------------

@app.get("/chat/stream/{request_id}")
async def stream_chat(request_id: str) -> StreamingResponse:
    req_data = _take_pending(request_id)

    async def generate() -> AsyncGenerator[str, None]:
        async for event in _run_agent_events(request_id, req_data):
            yield _sse_frame(event)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control":    "no-cache",
            "X-Accel-Buffering": "no",
            "Connection":       "keep-alive",
        },
    )


# ---------------------------------------------------------------------------
# GET /chat/status/{request_id}  — polling (alternativa a SSE)
# ---------------------------------------------------------------------------

async def _accumulate_events(request_id: str, req_data: dict) -> None:
    """Consume la corrida del agente en background y la vuelca al store de estado."""

    async def consume() -> None:
        async for event in _run_agent_events(request_id, req_data):
            if event["type"] == "step":
                request_states.add_step(request_id, event["data"])
            elif event["type"] == "final":
                request_states.set_final(request_id, event["data"])
            elif event["type"] == "error":
                request_states.set_error(
                    request_id, event["data"]["code"], event["data"]["message"]
                )

    try:
        # Sin este timeout, una corrida colgada deja al cliente polleando para
        # siempre: a diferencia de SSE, acá no hay desconexión que la cancele.
        await asyncio.wait_for(consume(), timeout=RUN_TIMEOUT_SECONDS)
    except asyncio.TimeoutError:
        logger.error("Timeout de corrida (request_id=%s) tras %ss", request_id, RUN_TIMEOUT_SECONDS)
        request_states.set_error(request_id, "RUN_TIMEOUT", RUN_TIMEOUT_MESSAGE)
    except asyncio.CancelledError:
        request_states.set_error(request_id, "CANCELLED", "La corrida fue cancelada.")
        raise
    except Exception as exc:
        logger.error("Error acumulando eventos (request_id=%s): %s", request_id, exc, exc_info=True)
        request_states.set_error(request_id, "ORCHESTRATOR_ERROR", ORCHESTRATOR_ERROR_MESSAGE)
    finally:
        request_states.finish(request_id)


@app.get("/chat/status/{request_id}", response_model=ChatStatusResponse)
async def chat_status(request_id: str) -> ChatStatusResponse:
    """Devuelve el estado acumulado de un request.

    El primer llamado reclama el request pendiente y dispara la corrida en
    background; los siguientes devuelven el snapshot acumulado. El cliente deja
    de pollear cuando `status` es `done` o `error`.
    """
    if not request_states.is_tracked(request_id):
        req_data = _take_pending(request_id)
        request_states.start(request_id)
        task = asyncio.create_task(_accumulate_events(request_id, req_data))
        _background_tasks.add(task)
        task.add_done_callback(_background_tasks.discard)
        logger.info("GET /chat/status → corrida iniciada request_id=%s", request_id)

    snapshot = request_states.snapshot(request_id)
    if snapshot is None:
        raise HTTPException(status_code=404, detail="request_id no encontrado o ya consumido")
    return ChatStatusResponse(**snapshot)


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@app.get("/health")
def health() -> dict:
    return {
        "status": "ok",
        "agent": APP_NAME,
        "pending_requests": len(pending_requests),
        "tracked_requests": len(request_states),
    }


# ---------------------------------------------------------------------------
if __name__ == "__main__":
    uvicorn.run(
        app,
        host=os.environ["HOST"],
        port=int(os.environ["PORT"]),
        log_level=os.environ["LOG_LEVEL"],
    )
